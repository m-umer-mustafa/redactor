import logging
import fitz
import docx

logger = logging.getLogger(__name__)

class Redactor:
    """
    Applies permanent redactions to physical files after human approval.
    """
    
    @staticmethod
    def _build_entity_type_map(approved_snippets: list[dict] | None) -> dict[str, str]:
        mapping = {}
        for snippet in approved_snippets or []:
            txt = (snippet or {}).get("text", "")
            ent = (snippet or {}).get("entity_type", "PII")
            if txt and txt not in mapping:
                mapping[txt] = ent
        return mapping

    @staticmethod
    def _replacement_for(entity_text: str, redaction_style: str, entity_type_map: dict[str, str]) -> str:
        if redaction_style == "████████":
            return "████████"
        if redaction_style == "[PII_TYPE]":
            return f"[{entity_type_map.get(entity_text, 'PII')}]"
        return "[REDACTED]"

    @staticmethod
    def apply_pdf_redactions(
        input_path: str,
        output_path: str,
        approved_entities: list[str],
        approved_snippets: list[dict] | None = None,
        redaction_style: str = "[REDACTED]",
    ) -> bool:
        """
        Uses PyMuPDF to draw permanent blackout boxes over matched PII strings.
        This actually removes the underlying font data from the PDF dictionary.
        """
        try:
            doc = fitz.open(input_path)
            # Remove exact duplicates from the list
            unique_entities = set(approved_entities)
            entity_type_map = Redactor._build_entity_type_map(approved_snippets)
            
            for page in doc:
                for entity_text in unique_entities:
                    # search_for returns a list of Rect objects where string is found
                    areas = page.search_for(entity_text)
                    for rect in areas:
                        replacement = Redactor._replacement_for(entity_text, redaction_style, entity_type_map)
                        # Add a redaction annotation (black color by default)
                        page.add_redact_annot(rect, text=replacement)
                        
                # Actually permanently apply the redactions to the page content
                page.apply_redactions()
                
            # Save the document, removing any invisible text/garbage underneath
            doc.save(output_path, garbage=3, deflate=True)
            return True
            
        except Exception as e:
            logger.error(f"Failed PDF redaction: {e}")
            return False

    @staticmethod
    def apply_docx_redactions(
        input_path: str,
        output_path: str,
        approved_entities: list[str],
        approved_snippets: list[dict] | None = None,
        redaction_style: str = "[REDACTED]",
    ) -> bool:
        """
        Iterates over DOCX paragraphs and tables, replacing matched PII with '[REDACTED]'.
        """
        try:
            doc = docx.Document(input_path)
            unique_entities = set(approved_entities)
            entity_type_map = Redactor._build_entity_type_map(approved_snippets)
            
            # Helper to replace text in runs to preserve some formatting
            def replace_text_in_runs(paragraphs):
                for p in paragraphs:
                    for entity_text in unique_entities:
                        if entity_text in p.text:
                            replacement = Redactor._replacement_for(entity_text, redaction_style, entity_type_map)
                            p.text = p.text.replace(entity_text, replacement)
                            
            # Process standard paragraphs
            replace_text_in_runs(doc.paragraphs)
            
            # Process table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_runs(cell.paragraphs)
                        
            doc.save(output_path)
            return True
            
        except Exception as e:
            logger.error(f"Failed DOCX redaction: {e}")
            return False

    @classmethod
    def redact(
        cls,
        input_path: str,
        output_path: str,
        approved_entities: list[str],
        approved_snippets: list[dict] | None = None,
        redaction_style: str = "[REDACTED]",
    ) -> bool:
        """Auto format detection for redaction output."""
        if not approved_entities:
            logger.warning("No entities approved for redaction. Proceeding with copy.")
            
        if input_path.lower().endswith('.pdf'):
            return cls.apply_pdf_redactions(
                input_path,
                output_path,
                approved_entities,
                approved_snippets=approved_snippets,
                redaction_style=redaction_style,
            )
        elif input_path.lower().endswith('.docx'):
            return cls.apply_docx_redactions(
                input_path,
                output_path,
                approved_entities,
                approved_snippets=approved_snippets,
                redaction_style=redaction_style,
            )
        else:
            raise ValueError("Unsupported format.")
