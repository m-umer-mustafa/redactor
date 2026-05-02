import os
import docx

def create_medical_record():
    doc = docx.Document()
    doc.add_heading('CONFIDENTIAL PATIENT ADMISSION RECORD', 0)
    doc.add_paragraph("Patient Name: Jonathan H. Abernathy")
    doc.add_paragraph("DOB: 12/04/1982 | Social Security Number: 293-11-9842")
    doc.add_paragraph("Contact Information:\nHome: (415) 882-9914\nEmail: j.abernathy_82@protonmail.com\nAddress: 8821 Westheimer Rd, Apt 4B, Houston, TX 77063")
    doc.add_heading('Clinical Notes', level=1)
    doc.add_paragraph(
        "Mr. Abernathy was admitted to St. Jude Medical Center at 14:00 hours complaining of "
        "acute chest pain. His emergency contact, Margaret Abernathy (Wife), can be reached at "
        "margaret.abernathy@lawfirm.com or 512-443-1929. The patient's primary care physician is "
        "Dr. Gregory House, operating out of Princeton-Plainsboro Teaching Hospital."
    )
    doc.save('demo_data/01_Medical_Record_Abernathy.docx')

def create_financial_audit():
    doc = docx.Document()
    doc.add_heading('Q3 Financial Audit & Disclosures', 0)
    doc.add_paragraph("Auditor: Ernst & Young LLP")
    doc.add_paragraph("Client: Global Payment Solutions Inc.")
    doc.add_heading('Discrepancy Report', level=1)
    doc.add_paragraph(
        "During the Q3 audit, several unauthorized wire transfers were detected originating from "
        "the corporate account. On August 14th, a sum of $45,000 was routed to an offshore account "
        "with IBAN: GB29BOSC60161331926819. The corresponding SWIFT code was BOFAUS3N. "
    )
    doc.add_paragraph(
        "Additionally, a corporate credit card registered to the CFO, Richard M. Sterling, was "
        "used for a $12,000 purchase at a luxury retailer in Dubai. The card number on file is "
        "4532 1123 9084 1128 (Visa) with an expiration of 12/28. Richard claims his wallet was "
        "stolen while staying at the Burj Al Arab hotel."
    )
    doc.save('demo_data/02_Financial_Audit_GPS.docx')

def create_HR_complaint():
    doc = docx.Document()
    doc.add_heading('INTERNAL HR COMPLAINT - TICKET #4491', 0)
    doc.add_heading('Complainant Information', level=1)
    doc.add_paragraph("Name: Sarah Jenkins\nDepartment: Logistics\nEmployee ID: 881-22A")
    doc.add_paragraph("Personal Phone: +1 206 555 0192\nPersonal Email: sj_logistics@yahoo.com")
    doc.add_heading('Incident Details', level=1)
    doc.add_paragraph(
        "I am writing this formal complaint regarding my manager, David R. Wallace. Over the past "
        "three weeks, David has continuously sent inappropriate text messages to my personal cell phone, "
        "which he obtained from my mandatory HR file. For example, on Tuesday, he texted me from his "
        "number (212) 867-5309 saying we should meet at the Marriott Hotel on 5th Avenue."
    )
    doc.add_paragraph(
        "I have forwarded screenshots of these conversations to my union representative, Michael Scott, "
        "who can be reached at m.scott@dundermifflin.com. I expect immediate action to be taken "
        "before I escalate this to the Equal Employment Opportunity Commission (EEOC) located at "
        "131 M Street, NE, Washington, DC 20507."
    )
    doc.save('demo_data/03_HR_Complaint_Jenkins.docx')

def create_raw_chat_log():
    doc = docx.Document()
    doc.add_heading('Exported Slack Transcription - #Project-Phoenix', 0)
    doc.add_paragraph("[10:14 AM] Alice L: Hey Bob, did you get the signed NDA from Acme Corp?")
    doc.add_paragraph("[10:16 AM] Bob Vance: Yeah, I have it right here. Vance Refrigeration forwarded it.")
    doc.add_paragraph("[10:17 AM] Alice L: Great. Send the invoice to their billing dept. Email is billing@acmecorp.com")
    doc.add_paragraph("[10:22 AM] Bob Vance: Done. Oh btw, my direct deposit bounced. Can you update my "
                      "routing number? It's 121000358 and the account is 998827361 at Wells Fargo.")
    doc.add_paragraph("[10:25 AM] Alice L: Please don't put banking info in Slack! Email that to HR. I'll delete this.")
    doc.add_paragraph("[10:26 AM] Bob Vance: Sorry! My bad. Call me later at 310-449-1120.")
    doc.save('demo_data/04_Slack_Chat_Log.docx')

def create_international_lease():
    doc = docx.Document()
    doc.add_heading('COMMERCIAL LEASE AGREEMENT', 0)
    doc.add_paragraph("Landlord: European Real Estate Holdings GMBH")
    doc.add_paragraph("Tenant: Samantha K. Wells")
    doc.add_heading('Premises', level=1)
    doc.add_paragraph(
        "The Landlord agrees to lease the commercial office space located at Alexanderplatz 1, "
        "10178 Berlin, Germany, to the Tenant. The Tenant's permanent US address on file is "
        "1422 Riverside Drive, Austin, Texas 78704."
    )
    doc.add_heading('Payments', level=1)
    doc.add_paragraph(
        "Monthly rent of €4,500 shall be wired to the Landlord's account. "
        "IBAN: DE89370400440532013000 at Deutsche Bank. "
        "Any late fees will be reported to international credit bureaus linked to the Tenant's "
        "passport number (P8891234) and US Social Security Number 000-11-2093."
    )
    doc.save('demo_data/05_International_Lease.docx')

def generate_all():
    os.makedirs("demo_data", exist_ok=True)
    create_medical_record()
    create_financial_audit()
    create_HR_complaint()
    create_raw_chat_log()
    create_international_lease()
    print("Successfully generated 5 diverse PII challenge documents in demo_data/")

if __name__ == "__main__":
    generate_all()
