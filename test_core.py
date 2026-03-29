from ml_engine import MLEngine
from document_parsers import DocumentParser
import os
from docx import Document

# Creating hardcore testing docs
test_doc_long = "long_test.docx"
test_doc_zero = "zero_test.docx"

doc_long = Document()
doc_long.add_paragraph("This is a massive document." * 5000) # Gigantic string
doc_long.add_paragraph("There is a secret person lying here: John Doe. Call him at 555-1234 or email john.doe@cybernetics.com.")
doc_long.save(test_doc_long)

doc_zero = Document()
doc_zero.add_paragraph("This document has absolutely zero sensitive information. It only contains the physics equations for gravity: F = G * m1 * m2 / r^2.")
doc_zero.save(test_doc_zero)

engine = MLEngine(use_transformer=False)

print("Testing Zero PII Doc...")
raw_text = DocumentParser.parse(test_doc_zero)
results = engine.analyze_text(raw_text)
assert len(results) == 0, "Failed: Should detect zero PII"
print("Zero PII Test Passed.")

print("Testing Long Document...")
raw_text_long = DocumentParser.parse(test_doc_long)
results_long = engine.analyze_text(raw_text_long)
snippets = engine.get_candidate_snippets(raw_text_long, results_long)
assert any(r.entity_type == "PERSON" for r in results_long), "Failed: Missing John Doe in long document."
print("Long Document Test Passed.")

print("All hardcore ML tests passed successfully.")
