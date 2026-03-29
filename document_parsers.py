import logging
import fitz  # PyMuPDF
import docx

logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Handles extracting plain text from different document formats.
    This simple extraction is sufficient for the MVP's NER pass.
    """
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extracts all text from a PDF, preserving order as best as possible."""
        try:
            doc = fitz.open(file_path)
            full_text = []
            for page in doc:
                text_page = page.get_text("text")
                full_text.append(text_page)
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return ""

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extracts text from paragraphs and tables in a DOCX."""
        try:
            doc = docx.Document(file_path)
            full_text = []
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables (often used in legal docs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
                            
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return ""

    @classmethod
    def parse(cls, file_path: str) -> str:
        """Auto-detect format and parse."""
        if file_path.lower().endswith('.pdf'):
            return cls.extract_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return cls.extract_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Must be .pdf or .docx")
